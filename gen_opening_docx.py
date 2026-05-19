"""
开题报告 DOCX 生成脚本
====================

生成《开题报告_文档.docx》。

内容特点：
1. 含“框架与大纲”部分，先明确各章节与子章节的写作范围；
2. 含四章完整正文，长度目标 8000-10000 字左右；
3. 含初步实验结果、表格、图片占位说明；
4. 第二个创新点（DAF/FDLP）按“方案已明确、实验待补”的开题口径撰写；
5. 若 figures/fig_XX.png 已存在，则自动插图；否则插入图位占位说明。

用法：
    /opt/anaconda3/bin/python3 gen_opening_docx.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


OUT_DOCX = "开题报告_文档.docx"
FIG_DIR = Path("figures")

META = {
    "title": "基于对比置信度的领域知识挟持与软引导投机解码与领域吸收飞轮研究",
    "author": "[姓名占位]",
    "advisor": "[导师姓名占位]",
    "school": "[学校 / 学院占位]",
    "major": "[专业占位]",
    "date": "2026 年 4 月",
}


def set_doc_style(doc: Document):
    """设置中文文档的基础字体。"""
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        st = doc.styles[name]
        st.font.name = "黑体"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_center_title(doc: Document, text: str, size=20, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(size)
    r.bold = bold


def add_para(doc: Document, text: str, first_line_chars=2):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.first_line_indent = Pt(first_line_chars * 12)
    fmt.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(12)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.4
    r = p.add_run("• " + text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(12)


def add_table(doc: Document, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def add_figure_placeholder(doc: Document, fig_name: str, caption: str):
    """若图片存在则插图，否则插入占位说明。"""
    fig_path = FIG_DIR / fig_name
    if fig_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(fig_path), width=Inches(6.2))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[图片占位：{fig_name}，请用 Nano Banana 生成后替换]")
        r.italic = True
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(caption)
    cr.font.name = "宋体"
    cr._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cr.font.size = Pt(10.5)


def main():
    doc = Document()
    set_doc_style(doc)

    # 封面
    doc.add_paragraph()
    add_center_title(doc, "博士学位论文开题报告", size=22)
    doc.add_paragraph()
    add_center_title(doc, META["title"], size=18)
    doc.add_paragraph()
    add_center_title(doc, f"汇报人：{META['author']}", size=14, bold=False)
    add_center_title(doc, f"指导教师：{META['advisor']}", size=14, bold=False)
    add_center_title(doc, f"{META['school']}  {META['major']}", size=14, bold=False)
    add_center_title(doc, META["date"], size=14, bold=False)

    doc.add_page_break()

    # 提纲
    doc.add_heading("一、框架与大纲", level=1)
    add_para(
        doc,
        "本开题报告围绕“在线引导的领域知识注入”和“离线可吸收的参数化沉淀”两条主线展开，"
        "整体目标是在不损害通用大模型综合能力的前提下，使其既能够在垂直领域获得可观的准确率提升，"
        "又能够在工程部署层面保留可接受的吞吐效率。报告共分为四章：第一章回答为什么要做这一问题、其工程价值在哪里；"
        "第二章梳理国内外研究现状并明确本文的研究空缺；第三章给出本文的研究目标、两个创新点及总体技术路线；"
        "第四章说明实施方案、初步实验结果、创新点二的可行性分析以及后续研究计划。"
    )
    add_bullet(doc, "第一章“选题依据与工程应用价值”：阐述 LLM 部署中的速度—能力矛盾、投机解码在垂直领域的失效现象，以及本文问题的现实需求。")
    add_bullet(doc, "第二章“国内外研究现状与发展动态”：从投机解码加速、领域知识注入、PEFT 选层与自适应 rank、机制可解释性四条线索进行综述，并归纳研究空白。")
    add_bullet(doc, "第三章“研究目标与研究内容”：明确总体目标、分目标以及两个创新点。创新点一为 DSSD 软引导投机解码，创新点二为 DAF 领域吸收飞轮。")
    add_bullet(doc, "第四章“实施方案及可行性分析”：说明三模型单卡内聚架构、已有 DSSD 初步结果、创新点二的实施路径与待补实验、进度安排与风险应对。")
    add_para(
        doc,
        "在写作顺序上，本文先用简明大纲统一论证逻辑，再进入完整正文展开。这样既满足开题报告“先定框架、再写内容”的要求，"
        "也便于后续直接从本稿演化为学院格式版 Word 文档。"
    )

    doc.add_page_break()

    # 第一章
    doc.add_heading("一、选题依据与工程应用价值", level=1)
    doc.add_heading("1.1 大模型部署的根本矛盾：通用模型强但慢，领域模型专但弱", level=2)
    add_para(
        doc,
        "近年来，大语言模型在通用问答、复杂推理、多轮对话、代码生成等任务上取得了突破性进展，"
        "参数规模从数十亿快速提升到数百亿甚至更高，模型的知识覆盖面、推理链长度以及指令服从能力显著增强。"
        "然而，模型能力提升的另一面，是推理成本和部署开销呈指数级上升。以 32B 量级指令模型为例，其单次前向需要调度海量参数，"
        "在真实服务环境中即使采用 bfloat16 精度和单卡 H200 部署，吞吐量也只有二十余 token 每秒。对于医疗辅助、法律检索、金融投顾、"
        "工业问答等需要快速交互的垂直场景而言，这一速度很难满足实时使用需求。"
    )
    add_para(
        doc,
        "与之对应，3B 左右的小模型虽然在显存占用和推理延迟上更具优势，但其通用推理能力、跨领域知识覆盖和长链推理稳定性明显弱于大模型。"
        "工业界常见做法是：保留大模型作为“综合验收官”，再通过垂直数据对小模型进行针对性微调，让小模型成为某一领域的“专家草稿器”。"
        "这种“大通用 + 小专家”的组合为效率与能力折中提供了自然路径，也正是投机解码能够发挥价值的场景基础。"
    )
    add_para(
        doc,
        "但问题在于，传统投机解码的理论假设更适合体量相近、分布相近的模型组合。当 Draft 与 Target 之间存在明显体量差，"
        "尤其是 Draft 经过垂直领域微调而 Target 保持通用状态时，Draft 提出的 token 往往正是具有高专业度的领域词汇，"
        "而这些词恰恰也是通用大模型概率分布中的弱项。结果是：正确的领域 token 被最强的通用模型大量拒绝，投机解码的接受率急剧下降，"
        "甚至比直接运行大模型更慢。换言之，垂直领域既是最需要投机解码的地方，也是传统投机解码最容易失效的地方。"
    )

    doc.add_heading("1.2 投机解码在垂直领域中的失效现象", level=2)
    add_para(
        doc,
        "为验证这一现象，本文选取 MedMCQA 的 Surgery 外科子集作为主实验领域，构建了 Target=Qwen2.5-32B-Instruct、"
        "Base=Qwen2.5-3B-Instruct、Draft=Qwen2.5-3B-Instruct-Surgery 的三模型体系。其中 Draft 由 Base 在外科数据上进行全参数微调得到，"
        "其目的不是取代 Target，而是在外科术语、病理概念和局部诊疗推理方面提供额外的领域先验。"
    )
    add_para(
        doc,
        "实验结果显示：Pure Target 在 Surgery 验证集上的准确率为 0.650；未经微调的 Base 仅有 0.494；"
        "而微调后的 Draft-Surgery 提升至 0.574，相比 Base 增加了 8.03 个百分点，说明 Draft 确实学到了真实可用的领域知识。"
        "然而，当它们被直接放入标准投机解码框架中时，standard SD 的接受率仅为 0.212，吞吐量约为 5.3 token/s，"
        "不仅没有获得加速，反而远低于 Pure Target 的 27.3 token/s。"
    )
    add_para(
        doc,
        "这一现象表明，问题并不在于 Draft“没有足够专业”，而在于当前的验收机制无法正确使用这种专业性。"
        "标准验收概率采用 min(1, P_target/P_draft) 的形式，当 3B Draft 在某个专业词上极其自信而 32B Target 因知识盲区给出较低概率时，"
        "该比值会迅速趋近于 0。于是，一个在领域上明明更正确的 token，反而因为与通用大模型分布不一致而被系统性压制。"
        "这类错误在外科、法律和代码等专业场景都具有普遍性。"
    )
    add_para(
        doc,
        "从工程视角看，这意味着如果继续沿用传统 SD，垂直领域部署就会陷入两难：要么完全依赖大模型，准确率尚可但成本过高；"
        "要么引入专家小模型，却因验收链路设计不当而得不到速度和质量上的双重收益。因此，重新设计“领域知识如何进入验收公式”，"
        "就不再是一个边缘微调问题，而是一个决定整个系统是否能落地的核心问题。"
    )

    add_figure_placeholder(
        doc,
        "fig_02.png",
        "图 1　三模型单卡内聚架构示意（建议用 fig_02，对应 prompt 见 nano_banana_prompts.md 第 2 项）",
    )

    doc.add_heading("1.3 工程应用价值", level=2)
    add_para(
        doc,
        "本文选题具有明确的工程应用价值。首先，在医疗场景中，问题往往涉及专业术语密集、因果链较长且交互延迟敏感，"
        "如术后并发症分析、影像报告辅助解释、围术期风险评估等。若每次推理都完全交给 30B 以上模型，成本与时延均难以接受；"
        "若仅由几亿至几十亿参数的小模型处理，则很容易在非典型问法、多轮上下文或复杂归纳时失稳。本文提出的 DSSD 机制，"
        "恰恰是让大模型保留主体控制权，而让小模型只在最关键的专业位置上进行补位。"
    )
    add_para(
        doc,
        "其次，在法律、金融、工业质检、企业知识库问答等场景中，用户更在意“专业词是否正确”“关键实体是否遗漏”“系统是否稳定可靠”。"
        "这类任务不需要模型在每一个 token 上都表现出专家倾向，而是需要在极少数但决定性的专业位置上做出正确选择。"
        "因此，与其一味增大全局模型或全量微调，不如设计一个能精准识别“哪里需要专家知识”的推理与训练闭环。"
    )
    add_para(
        doc,
        "进一步地，创新点二 DAF 的工程意义在于：如果在线引导的领域信号能够被离线沉淀到 Target 的最小 PEFT 增量中，"
        "则未来部署时就不必每次都同时挂载 Draft 和 Base。换言之，第一点回答“领域知识如何在解码时被借用”，"
        "第二点回答“这种借用能否被固化为参数能力”。一旦这一闭环成立，系统可从‘运行时多模型协同’逐步过渡到‘训练时吸收、运行时单模型推理’，"
        "其部署价值远高于单纯的学术涨点。"
    )

    doc.add_heading("1.4 选题切入点的合理性", level=2)
    add_para(
        doc,
        "本文的切入点建立在两个已被实验证明的前提之上。其一，Base 与 Draft 的差值信号确实能够刻画“领域微调带来的概率增量”，"
        "因为二者同体量、同架构，唯一系统性差异就是领域数据的吸收；其二，Target 的熵信号能够刻画“通用大模型对当前位置是否真正不确定”，"
        "因为在 Surgery 与 Anatomy 等领域盲区中，错题熵显著高于对题，而在 Pharmacology 这类 Target 本已较熟悉的子领域中，熵差异并不显著。"
        "这意味着我们不是凭直觉构造门控，而是在用两个可观测、可解释、且已被实证支撑的信号决定何时该让领域知识介入。"
    )
    add_para(
        doc,
        "更重要的是，第一点与第二点之间并非简单并列，而是存在天然递进关系。第一点先证明“领域知识在解码时确实有效”；"
        "第二点再追问“这种有效性是否能沉淀为权重中的最小增量”。前者解决存在性问题，后者解决可吸收性与可部署性问题。"
        "对于博士开题而言，这种从理论机制、系统工程到参数化沉淀的递进链条具有较强的完整性，也更容易形成后续多篇论文的持续研究主线。"
    )

    # 第二章
    doc.add_heading("二、国内外研究现状与发展动态", level=1)
    doc.add_heading("2.1 投机解码加速方向的研究现状", level=2)
    add_para(
        doc,
        "投机解码是近年来大模型推理加速领域最活跃的研究方向之一。早期工作主要围绕“让小模型先提案，再让大模型并行验收”这一基本框架展开，"
        "核心目标是在保持输出分布尽量不变的前提下减少 Target 的串行调用次数。随后，一系列后续工作开始探索如何提升草稿质量或提高单次 Target 调用的有效产出，"
        "例如使用多 head 直接并行预测多个未来 token 的 Medusa、利用层间退出或结构重组增强草稿质量的 Eagle、以及通过树状分支并行验证扩大单次覆盖范围的 SpecInfer 等。"
    )
    add_para(
        doc,
        "这些工作在通用场景下验证了投机解码可以达到较高的理论加速比，但也普遍依赖两个条件：一是 Draft 与 Target 的分布差异不能太大；"
        "二是草稿错误不能高度集中在关键位置。对于医学、法律等垂直领域，情况恰恰相反：正确与错误往往不再平均分散，而是集中在专业词和关键实体上。"
        "因此，仅从“怎么让草稿更快”或“怎么让验收更并行”出发，尚不足以解决领域场景下的根本问题。"
    )
    add_para(
        doc,
        "此外，越来越多研究开始尝试把投机解码与训练联动，例如为 Draft 设计接受率导向的训练损失，或者让草稿模型在在线运行中不断适应 Target 的偏好。"
        "这说明学界已经注意到“验收事件本身包含学习信号”，但现有工作大多聚焦于 Draft 训练，较少讨论如何利用这些事件反向优化 Target 的 PEFT 放置策略。"
    )

    doc.add_heading("2.2 领域知识注入与在线 steering 的研究现状", level=2)
    add_para(
        doc,
        "另一条与本文紧密相关的主线，是领域知识注入与推理期 steering。相关研究通常尝试在概率域、logit 域或激活域中向目标模型注入某种外部信号，"
        "例如 Product of Experts 通过乘积形式将多个分布耦合，Steering Vector 通过特定方向的残差修改实现行为偏置，"
        "还有一些方法直接在 decoder 端进行融合，用以提升某些受限属性或领域偏好。"
    )
    add_para(
        doc,
        "这些方法的共同优点是无需重新训练大模型，具有良好的即插即用性；但它们也存在明显局限：第一，许多方法在全局范围内持续注入，"
        "容易在本不需要外部知识干预的位置引入噪声；第二，注入信号通常来自单一来源，例如仅依赖专家模型置信度或单一 steering 向量，"
        "缺少“Target 自身是否真的需要帮助”的内部判据；第三，这类方法往往停留在在线使用层面，并未进一步回答“在线引导能否被离线吸收为参数能力”的问题。"
    )
    add_para(
        doc,
        "本文的 DSSD 与该方向的差异在于：我们不直接把 Draft 当作外部专家分布全局融合，而是先引入同体量的 Base 作为常识对照组，"
        "提取 Draft 相对 Base 的超额置信度 ΔP。与此同时，我们还引入 Target 的 Shannon 熵作为内部路由信号，使领域知识仅在‘专家自信且大模型不确定’的步骤被放大。"
        "这种双信号门控比全程注入更稀疏、更精准，也更符合垂直领域场景中的真实需求。"
    )

    doc.add_heading("2.3 PEFT 选层与自适应 rank 分配的研究现状", level=2)
    add_para(
        doc,
        "在参数高效微调方面，LoRA、QLoRA、Adapter 及其变体已经成为大模型领域适配的主流方案。随着应用深入，研究关注点逐步从“要不要用 PEFT”转向"
        "“在同样预算下，LoRA 应该加在哪里、rank 应该如何分配”。代表性方法包括基于梯度或近似二阶信息做自适应 rank 分配的 AdaLoRA，"
        "基于重要性排序或激活统计进行模块筛选的 Flexora、Act-LoRA、GoRA、IGU-LoRA 等。"
    )
    add_para(
        doc,
        "这些工作为本文提供了重要启发：LoRA 放置并非越多越好，而是需要依据任务敏感层做稀疏投放。但与此同时，"
        "现有工作大多采用 task-level loss、全数据集梯度统计、或层级激活幅值作为重要性依据，其监督锚点仍然是宏观的。"
        "它们能够回答‘哪些层对总体损失下降有贡献’，却较难回答‘哪些层与某一类关键 token 决策改写直接相关’。"
    )
    add_para(
        doc,
        "本文创新点二的核心不同就在于，我们不是从所有 token 上平均统计重要性，而是只在由 DSSD 触发的 flip 事件上统计敏感度。"
        "换言之，FDLP 并不是再造一种新的 PEFT 算法，而是为 PEFT placement 引入一种解码事件驱动的监督锚点。"
        "如果这一锚点相比 entropy、一般分歧或 all-token 统计具有明显优势，那么就说明投机解码中的事件日志本身携带了此前未被利用的结构化训练价值。"
    )

    doc.add_heading("2.4 机制可解释性工具的发展动态", level=2)
    add_para(
        doc,
        "近年来，Activation Patching、Causal Tracing、Representation Editing 等机制可解释性方法被广泛用于分析 Transformer 内部的因果链条。"
        "这些工具能够回答“某层某个表示是否真正处在某一行为生成的关键路径上”，因此非常适合被用作 PEFT 选层的事后因果验收。"
        "不过，需要强调的是，这类工具本身并不是本文的创新点，它们更像是验证工具箱。"
    )
    add_para(
        doc,
        "本文计划在 DAF 中将 patching 用作验证步骤：如果 FDLP 选出的热点层同时也是 patching 后最能提升 flip 目标 token 概率的层，"
        "则可以较有力地说明这些层不是偶然相关，而是真正位于领域知识改写的因果杠杆点上。"
        "这对于博士论文而言非常重要，因为它使方法论不止停留在工程经验，而具备一定的机制解释支撑。"
    )

    doc.add_heading("2.5 现状总结与本文定位", level=2)
    add_para(
        doc,
        "综合以上综述可以发现，当前研究虽然已经分别在投机解码、在线 steering、PEFT placement 与机制验证上取得了长足进展，"
        "但仍缺少一条真正打通“解码期事件”与“训练期放置”的闭环主线。传统投机解码关注速度，领域注入关注在线修正，"
        "PEFT 选层关注参数预算，而机制解释关注因果验证。这四条线索往往各自独立，尚未形成一个统一框架。"
    )
    add_para(
        doc,
        "本文的定位正是在这一交叉空白处：第一，提出 DSSD，在投机解码验收公式内部显式引入领域知识探针 ΔP 与双信号门控，解决“领域知识如何被借用”；"
        "第二，提出 DAF/FDLP，把 DSSD 产生的 flip 事件同时用作 LoRA 选层监督和飞轮停止准则，解决“被借用的知识能否被吸收”。"
        "因此，本文并不是对已有某条线索的简单微调，而是试图建立一个从在线推理到离线训练、从统计提升到机制解释的系统性研究框架。"
    )

    add_figure_placeholder(
        doc,
        "fig_01.png",
        "图 2　DSSD 与 DAF 的总体技术路线（建议用 fig_01，对应 prompt 见 nano_banana_prompts.md 第 1 项）",
    )

    # 第三章
    doc.add_heading("三、研究目标与研究内容", level=1)
    doc.add_heading("3.1 总体研究目标", level=2)
    add_para(
        doc,
        "本文总体目标可概括为：构建一个面向垂直领域的大模型推理与吸收双闭环框架，在不修改或尽量少修改大模型主体权重的前提下，"
        "实现领域准确率的显著提升，并在最终部署阶段尽量恢复单模型推理的吞吐效率。具体而言，本文希望首先在零训练前提下，通过推理期的软引导投机解码，"
        "使 Target 在领域任务上的准确率显著优于自身基线；随后，再利用推理期生成的事件日志，指导最小化 PEFT 增量，使这些在线收益逐步沉淀为离线能力。"
    )
    add_para(
        doc,
        "从可量化指标看，创新点一已经给出了明确目标并取得初步结果：在 MedMCQA Surgery 子集上，使 32B Target 的准确率从 0.650 提升至 0.700，"
        "即提升 5 个百分点；同时，在一组可选工作点上实现较优的速度—质量帕累托前沿。创新点二则以“吸收后单独运行 Target 仍能接近 DSSD 在线增益，"
        "且 tps 恢复到接近 pure_target”为目标，重点验证领域知识的可吸收性与飞轮机制的存在性。"
    )

    doc.add_heading("3.2 创新点一：DSSD 软引导投机解码", level=2)
    add_para(
        doc,
        "创新点一的基本思想，是将微调 Draft 所体现的领域优势，与未微调 Base 所体现的通用常识进行显式对比，"
        "并把二者之差作为领域知识探针信号。其数学形式可写为 ΔP(x)=P_draft(x)-P_base(x)。与直接使用 P_draft 不同，"
        "这种设计能够自动消去大量通用高频词上的共同概率质量，只保留因领域微调而新增的那一部分偏好，因此更适合刻画“真正的领域知识位置”。"
    )
    add_para(
        doc,
        "在验收阶段，本文围绕标准 SD 的 P_accept=min(1, P_target/P_draft) 构造了一系列软引导策略。"
        "其中，C1 在比值域上引入 α·ΔP 补偿；C3 在概率域上进行量纲一致的 subsidy；C4 用 Draft 自信度优势 S_t 做步级门控；"
        "C5 用 Target 熵 H_t/H_max 表示内部不确定性；C6 进一步把两者相乘形成双信号 AND 门；C8 再将步级门控推进到 token 级门控，"
        "使门控粒度与验收 token 完全对齐。整套策略从固定强度、单信号、步级门控，逐步演化到双信号、token 级稀疏激活，"
        "构成了一个完整且可消融的设计谱系。"
    )
    add_para(
        doc,
        "现有实验表明，这一谱系不仅能涨点，而且能解释为什么涨点。C1 在 α=0.10 时即可将准确率提升到 0.690，说明在大体量差场景下，"
        "比值域加法确实具有良好的鲁棒性；C5 在 λ=5 时达到 0.685，证明 Target 熵可以作为较有效的不确定性路由信号；"
        "C6 在 λ=50 时达到 0.690，同时 tps 提升到 11.5，说明双信号门控在保持高准确率的同时改善了接受率；"
        "C8 则在 λ=20 时进一步达到 0.700 的全局最高准确率，验证了“门控粒度与验收 token 对齐”这一设计原则。"
    )

    add_table(
        doc,
        ["策略", "最优超参数", "acc", "tps", "说明"],
        [
            ["pure_target", "—", "0.650", "27.3", "基线"],
            ["standard_sd", "—", "0.650", "5.3", "无加速，接受率低"],
            ["C1", "α=0.10", "0.690", "6.9", "最早验证 ΔP 有效"],
            ["C5", "λ=5", "0.685", "8.2", "熵权路由"],
            ["C6", "λ=50", "0.690", "11.5", "Pareto 突破"],
            ["C8", "λ=20", "0.700", "12.2", "当前全局最优"],
            ["C9", "λ=100", "0.690", "15.7", "tps 最优工作点"],
        ],
    )

    add_figure_placeholder(
        doc,
        "fig_03.png",
        "图 3　C1–C9 软引导策略演进树（建议用 fig_03，对应 prompt 见 nano_banana_prompts.md 第 3 项）",
    )
    add_figure_placeholder(
        doc,
        "fig_04.png",
        "图 4　C6 双信号联合门控架构（建议用 fig_04，对应 prompt 见 nano_banana_prompts.md 第 4 项）",
    )

    doc.add_heading("3.3 创新点二：DAF 领域吸收飞轮", level=2)
    add_para(
        doc,
        "如果说创新点一解决的是“在线怎么借”，那么创新点二要解决的是“借来的知识能不能沉淀”。"
        "基于这一目标，本文提出 Domain Absorption Flywheel，即把 DSSD 在解码过程中产生的 flip 事件赋予双重角色："
        "一方面，flip 事件代表 Target 的原始 argmax 被 Draft 提案改写，这说明领域知识在该处真实改变了决策边界，因此它可以作为稀疏监督锚点；"
        "另一方面，flip rate 的变化又能直接反映当前轮 LoRA 训练是否已经吸收了一部分原本需要在线引导的知识，因此它也可以作为飞轮收敛的可观测量。"
    )
    add_para(
        doc,
        "在具体实现上，DAF 的每一轮包含四步：首先用固定策略（主线计划使用 C9）在训练集上运行 DSSD，记录所有 flip 事件及其 prefix、"
        "原始 token A_t、最终 token B_t、ΔP、熵等字段；其次，在这些 flip 样本上计算 Flip-driven LoRA Placement 的层敏感度分数，"
        "得到最适合挂载 LoRA 的 Top-K 层；再次，在这些层上进行小规模 LoRA 训练，并可选择将适配器 merge 回 Target；最后，用更新后的 Target 重新运行解码，"
        "观察 flip rate 是否显著下降。若跨轮变化低于阈值，或通用域退化超过 1%，则停止飞轮。"
    )
    add_para(
        doc,
        "这一设计最关键的地方在于，它把“解码期观察到的结构化事件”与“训练期参数预算的投放位置”严格绑定起来。"
        "与一般的自适应 rank 方法不同，FDLP 并不在所有 token 上平均统计重要性，而是只在 F_t=1 的稀疏事件上计算梯度敏感度；"
        "与普通迭代 SFT 不同，DAF 也不是简单地一轮轮做微调，而是要求每一轮必须带来 flip 分布的变化，并接受 entropy / disagreement / all-token 事件作为严格对照。"
        "如果只有 flip 驱动的飞轮能够稳定收敛，那么就能够较强地支持“投机解码事件具有独特训练价值”的主张。"
    )

    add_figure_placeholder(
        doc,
        "fig_05.png",
        "图 5　DAF 领域吸收飞轮闭环（建议用 fig_05，对应 prompt 见 nano_banana_prompts.md 第 5 项）",
    )
    add_figure_placeholder(
        doc,
        "fig_06.png",
        "图 6　FDLP 选层算法流程（建议用 fig_06，对应 prompt 见 nano_banana_prompts.md 第 6 项）",
    )
    add_figure_placeholder(
        doc,
        "fig_07.png",
        "图 7　flip 事件的双重角色（建议用 fig_07，对应 prompt 见 nano_banana_prompts.md 第 7 项）",
    )

    doc.add_heading("3.4 关键科学问题与研究路线图", level=2)
    add_para(
        doc,
        "围绕上述两个创新点，本文要回答四个关键科学问题：第一，为什么 Draft−Base 的对比信号能比单独的 Draft 置信度更准确地识别领域 token；"
        "第二，为什么需要同时引入 Draft 自信度和 Target 熵，单一信号为何不足；第三，为什么 flip 事件比 entropy 或一般分歧更适合作为 PEFT 放置的监督锚点；"
        "第四，在线推理中的局部事件，是否真的能够被离线最小参数更新吸收并保留下来。"
    )
    add_para(
        doc,
        "对应地，本文研究路线也呈现出清晰的阶段化结构：第一阶段完成 DSSD 各策略的系统消融，明确最优工作点与关键机制；第二阶段完成 Target 熵探针、"
        "Draft 延续词探针、硬替换因果验证等解释性实验，增强创新点一的理论与实验支撑；第三阶段实现 flip 日志系统、FDLP 选层与单轮 LoRA 对照；"
        "第四阶段进一步扩展为完整的 DAF 飞轮与事件对照实验；第五阶段再考虑 vLLM 集成、树状投机解码、跨领域泛化等扩展问题。"
    )

    # 第四章
    doc.add_heading("四、实施方案及可行性分析", level=1)
    doc.add_heading("4.1 硬件平台与系统实现方案", level=2)
    add_para(
        doc,
        "本文实验平台采用单卡 H200（141GB HBM3e）作为核心硬件，所有模型统一运行在 cuda:0 上，严格避免任何跨卡通信。"
        "Target 使用 Qwen2.5-32B-Instruct，Base 使用 Qwen2.5-3B-Instruct，Draft 使用在 MedMCQA Surgery 数据上微调得到的 3B 模型。"
        "全局精度统一为 bfloat16。为了保证实验公平性和系统可控性，本文禁用高阶 pipeline 与 model.generate 接口，全部基于 model.forward 手写 auto-regressive 解码循环。"
    )
    add_para(
        doc,
        "在缓存管理方面，本文采用 StaticCache 进行 KV Buffer 预分配，以支持前缀共享、Lazy Evaluation 与 Shadow Sync 模式。"
        "其中，Base 在多步提案阶段跟随 Draft 一起向前推进，但仅在候选 token 位置按需计算 LM Head，从而降低无效 softmax 的开销。"
        "这一工程设计保证了三模型共驻单卡在显存上可行，也为后续采集完整的解码遥测信息提供了统一执行引擎。"
    )

    doc.add_heading("4.2 创新点一的初步实验结果", level=2)
    add_para(
        doc,
        "目前，创新点一已经完成领域选择、Draft 微调、主策略消融、熵探针和部分因果验证实验。"
        "在领域选择上，本文比较了 MedQA-USMLE、JEC-QA、CMB-Exam 与 MedMCQA Surgery 等多个候选方向，"
        "最终选择 Target 仍有明显提升空间、且训练数据足够丰富的 Surgery 作为主战场。"
        "在 Draft 微调上，基于 22K 左右 Surgery 训练样本与 25% Alpaca 格式锚点，最终最优 checkpoint 将 3B Base 的准确率从 0.494 提升到 0.574。"
    )
    add_para(
        doc,
        "在软引导策略方面，C1、C5、C6、C8、C9 已形成清晰的帕累托前沿。C1 证明了比值域补偿的有效性，"
        "C5 证明了 Target 熵的路由价值，C6 证明双信号 AND 门可同时提升准确率和接受率，C8 则通过 token 级门控把总体准确率进一步推进到 0.700。"
        "换言之，创新点一已经不仅是一个概念设计，而是一个具备完整实验支撑的研究闭环。"
    )

    add_figure_placeholder(
        doc,
        "fig_08.png",
        "图 8　DSSD 的 acc-tps 帕累托前沿（建议用 fig_08，对应 prompt 见 nano_banana_prompts.md 第 8 项）",
    )

    add_para(
        doc,
        "此外，Target 熵探针结果也提供了强支撑。三科目实验表明，Surgery 错题平均熵 0.456、对题平均熵 0.387，差异显著（p=0.0008）；"
        "Anatomy 也表现出显著差异（0.409 vs 0.309，p=0.0004）；而 Pharmacology 的差异不显著（p=0.1218）。"
        "这说明 Target 熵确实可以作为领域盲区的探测器，而不是对所有子领域都一视同仁地高涨。进一步地，ΔP>0.05 的领域词位置熵为 0.611，"
        "相比通用词位置的 0.390 高出 57%，从词级别直接支持了“高熵 + 高 ΔP 位置就是最值得注入的地方”。"
    )

    add_table(
        doc,
        ["科目", "Target acc", "对题平均熵", "错题平均熵", "p 值", "结论"],
        [
            ["Surgery", "59%", "0.387", "0.456", "0.0008", "显著"],
            ["Pharmacology", "76%", "0.321", "0.350", "0.1218", "不显著"],
            ["Anatomy", "74%", "0.309", "0.409", "0.0004", "显著"],
        ],
    )

    add_figure_placeholder(
        doc,
        "fig_09.png",
        "图 9　三科目 Target 熵分布对比（建议用 fig_09，对应 prompt 见 nano_banana_prompts.md 第 9 项）",
    )

    doc.add_heading("4.3 创新点二的实施步骤与待补实验清单", level=2)
    add_para(
        doc,
        "与创新点一相比，创新点二目前仍处于“方案设计已明确、实验实施即将开始”的阶段，这也是开题报告阶段较为合理的研究状态。"
        "从工程增量上看，创新点二并不需要推翻现有系统，而是在已有 DSSD 管线上做三类新增工作：其一，扩展 TelemetryLogger，记录 A_t、B_t、F_t、"
        "prefix_ids、delta_P、H_t 等字段，形成可复用的 flip 事件日志；其二，实现 FDLP 打分脚本，对 flip 样本做只读反向传播，得到层敏感度排序；"
        "其三，把排序结果自动转化为 LoRA 配置，完成单轮和多轮实验。"
    )
    add_para(
        doc,
        "考虑到博士开题报告强调“问题是否值得做、路线是否可行”，而不要求所有实验在开题前全部完成，因此本文建议在开题答辩前先补齐最小实验集合。"
        "只要这组最小实验能够跑通，就足以证明第二个创新点不是空洞设想，而是具有明确实施入口和可证伪路径的研究方案。"
    )

    add_table(
        doc,
        ["优先级", "建议补充实验", "目的", "预期输出"],
        [
            ["P0", "flip 事件日志采集（Round 0）", "验证 F_t 定义可稳定记录", "flip rate、事件样本数、样例 JSON"],
            ["P0", "FDLP 单轮选层热图", "验证热点层是否集中而非随机", "layer score 排序图"],
            ["P0", "Top-K LoRA vs 全层 LoRA", "验证稀疏选层是否具备效果", "一组 acc 对照表"],
            ["P1", "entropy / disagreement 选层对照", "保护创新边界", "事件必要性对照结果"],
            ["P1", "MMLU 小样本守护集", "验证通用能力未明显退化", "退化曲线或表格"],
        ],
    )

    doc.add_heading("4.4 工程可行性与科学可行性分析", level=2)
    add_para(
        doc,
        "从工程可行性看，本文方案具有较好的渐进式实现特征。创新点一已经在当前 H200 单卡环境上稳定运行，"
        "这意味着三模型内聚、KV Cache 管理、遥测日志采集和领域数据处理均已打通。创新点二只需在此基础上增加日志字段、只读反向传播和 LoRA 配置生成，"
        "工程复杂度远低于重新搭建一套训练—推理系统。单轮 32B LoRA 训练虽然有一定显存与时间压力，但在 H200 上结合 bfloat16、checkpointing 和小 batch 设定仍具备可操作性。"
    )
    add_para(
        doc,
        "从科学可行性看，创新点二也建立在较强的前提支撑上。首先，第一点已经证明 flip 事件并不是随机噪声，而是与领域 token、Target 高熵位置和 Draft 专业续写高度相关；"
        "其次，第一点的最佳策略并非在所有位置都注入，而是稀疏地修改关键决策边界，这使得 flip 事件天然适合作为稀疏监督锚点；"
        "再次，现有 PEFT 与机制解释工作已经提供了成熟工具箱，使得“选层—微调—验证”的技术路线不存在根本方法障碍。"
        "因此，创新点二最大的挑战不是能否做出来，而是最终效果能否显著优于强基线；而这正是开题后重点要用实验回答的问题。"
    )

    doc.add_heading("4.5 风险识别与应对策略", level=2)
    add_para(
        doc,
        "本文预判至少存在四类风险。第一，飞轮可能在第一轮就收敛，即 K*=1。这并不意味着方案失败，而更可能说明 Surgery 子集上的可吸收信号并不需要多轮迭代；"
        "此时应将贡献表述收缩为“FDLP 单轮 + 收敛判据”。第二，飞轮可能把 Target 推向 Draft，导致通用域能力退化。对此需要使用 MMLU 或 GSM8K 小样本守护集，"
        "一旦退化超过 1%，立即回退到上一轮 Target。第三，热点层可能跨轮漂移，导致难以形成稳定 placement。对此应显式报告 Top-K 层 Jaccard 相似度，并把‘漂移’本身视作可发表现象。"
        "第四，entropy 或 disagreement 事件可能在飞轮中表现接近 flip。若出现这种情况，就必须诚实收缩创新性主张，避免过度包装。"
    )

    doc.add_heading("4.6 进度安排与阶段目标", level=2)
    add_para(
        doc,
        "根据当前研究基础与博士培养周期，本文后续计划分为四个阶段推进。第一阶段（2026 年 5 月至 8 月）完成 DSSD 各策略收尾、通用能力测试和论文投稿准备；"
        "第二阶段（2026 年 7 月至 10 月）实现 flip 日志系统、FDLP 单轮选层和一轮 LoRA 对照；第三阶段（2026 年 10 月至 2027 年 2 月）开展 DAF 多轮飞轮、"
        "事件对照和守护集实验；第四阶段（2027 年 3 月至 6 月）进行系统优化、论文撰写与答辩准备。整个过程中，每一阶段都对应可独立汇报的里程碑成果，"
        "确保研究不会因为某一个子假设不成立而整体失去方向。"
    )

    add_figure_placeholder(
        doc,
        "fig_10.png",
        "图 10　后续研究进度甘特图（建议用 fig_10，对应 prompt 见 nano_banana_prompts.md 第 10 项）",
    )

    # 结尾
    doc.add_heading("结语", level=1)
    add_para(
        doc,
        "总体而言，本文选题具有清晰的问题牵引、扎实的初步实验基础和明确的后续扩展路线。创新点一已经完成从理论动机、策略设计到实证验证的闭环，"
        "证明了在大体量差场景下，领域知识并非不能进入投机解码，而是需要一个更符合垂直场景的软引导验收机制。创新点二则进一步把这一机制推进到参数化吸收层面，"
        "尝试回答‘在线借来的知识，能否最终变成 Target 自己的能力’这一更具博士论文价值的问题。若研究顺利推进，本文不仅有望形成关于领域引导投机解码的完整理论与系统框架，"
        "也可能为未来垂直领域大模型的低成本部署提供一条新的工程路径。"
    )

    # 附录：图片清单
    doc.add_heading("附录：建议配图清单", level=1)
    for idx, desc in [
        ("fig_01", "总体技术路线图"),
        ("fig_02", "三模型单卡内聚架构"),
        ("fig_03", "C1–C9 策略演进树"),
        ("fig_04", "C6 双信号门控架构"),
        ("fig_05", "DAF 飞轮闭环"),
        ("fig_06", "FDLP 选层算法"),
        ("fig_07", "flip 事件双重角色"),
        ("fig_08", "Pareto 前沿热图"),
        ("fig_09", "三科熵分布对比图"),
        ("fig_10", "研究进度甘特图"),
    ]:
        add_bullet(doc, f"{idx}：{desc}；完整生成 prompt 见 `nano_banana_prompts.md`。")

    # 保存并统计文本长度
    doc.save(OUT_DOCX)
    text_len = sum(len(p.text) for p in doc.paragraphs)
    print(f"[OK] 已生成 {OUT_DOCX}")
    print(f"[INFO] 文本总长度（按段落字符计）≈ {text_len} 字")


if __name__ == "__main__":
    main()
